
import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from datetime import date
from io import BytesIO
import zipfile

st.title("Multi-Bureau Credit Dispute Letter Generator")

st.write("Upload a credit report PDF and generate dispute letters for Experian, Equifax, and TransUnion.")

uploaded_file = st.file_uploader("Upload your credit report (PDF)", type="pdf")

if uploaded_file:
    # Load and extract text
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    full_text = ""
    for page in doc:
        full_text += page.get_text()
    doc.close()

    # Simulated dispute flagging logic
    flagged_items = []
    lines = full_text.split("\n")
    for i, line in enumerate(lines):
        if "Collection account" in line or "charged off" in line:
            context = "\n".join(lines[i-3:i+3])
            flagged_items.append(context)

    if not flagged_items:
        st.warning("No obvious negative items found (collections or charge-offs).")
    else:
        st.success(f"Found {len(flagged_items)} potential dispute item(s).")

        # Simulated assignment to bureaus
        disputes_by_bureau = {
            "Experian": [
                {"account": "AFFIRM INC", "reason": "Charged off $1,469. Not my account or unverifiable debt."},
                {"account": "EXXONMOBIL/CBNA", "reason": "Charged off $422. I request validation or removal."}
            ],
            "Equifax": [
                {"account": "POSSIBLE FINANCIAL INC", "reason": "Charged off $155. Inaccurate and outdated information."}
            ],
            "TransUnion": [
                {"account": "NATIONAL CREDIT ADJUST", "reason": "Duplicate collection account reported. Please remove."}
            ]
        }

        bureau_addresses = {
            "Experian": ["Experian", "P.O. Box 4500", "Allen, TX 75013"],
            "Equifax": ["Equifax", "P.O. Box 740256", "Atlanta, GA 30374"],
            "TransUnion": ["TransUnion", "P.O. Box 2000", "Chester, PA 19016"]
        }

        user_name = "TREVIN COOKS"
        user_address = "103 C V JACKSON DR\nNEW IBERIA, LA 70560-4420"
        today = date.today().strftime("%B %d, %Y")

        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED) as zip_file:
            for bureau, items in disputes_by_bureau.items():
                doc = Document()
                doc.add_paragraph(user_name)
                for line in user_address.split("\n"):
                    doc.add_paragraph(line)
                doc.add_paragraph(today)
                doc.add_paragraph("")
                for line in bureau_addresses[bureau]:
                    doc.add_paragraph(line)
                doc.add_paragraph("")
                doc.add_paragraph("RE: Request for Investigation under FCRA §609 and §611")
                doc.add_paragraph("To Whom It May Concern,")
                doc.add_paragraph("")
                doc.add_paragraph("I am writing to formally dispute the accuracy of the following items listed on my credit report. In accordance with the Fair Credit Reporting Act (§609 and §611), I request that you investigate and remove the following inaccurate or unverifiable entries:")
                for item in items:
                    doc.add_paragraph(f"• Account: {item['account']}\n   - Reason: {item['reason']}")
                doc.add_paragraph("")
                doc.add_paragraph("Please provide a copy of the results of your investigation. I have enclosed a copy of my identification and proof of address for verification purposes.")
                doc.add_paragraph("")
                doc.add_paragraph("Sincerely,")
                doc.add_paragraph(user_name)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                zip_file.writestr(f"{bureau}_Dispute_Letter.docx", buffer.read())

        zip_buffer.seek(0)
        st.download_button(
            label="Download All Dispute Letters (ZIP)",
            data=zip_buffer,
            file_name="Dispute_Letters_All_Bureaus.zip",
            mime="application/zip"
        )
